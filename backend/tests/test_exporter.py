# backend/tests/test_exporter.py
# Tests for PDF and DOCX export service
# Related: app/services/exporter.py, app/models/schemas.py

import asyncio
from pathlib import Path

import pytest

from app.models.schemas import (
    AggregatedReport,
    ConfidenceNote,
    Deadlines,
    EstimatedValue,
    EvaluationCriterion,
    LotInfo,
    ProcuringOrganization,
    QAEvaluation,
    QualificationRequirements,
    SourceDocument,
)
from app.services.exporter import (
    _format_date,
    _format_value,
    _qa_color,
    _severity_color,
    export_docx,
    export_pdf,
)


# ── Fixtures ───────────────────────────────────────────────────────────────────


@pytest.fixture
def full_report() -> AggregatedReport:
    """Report with all fields populated."""
    return AggregatedReport(
        project_summary="Vilniaus miesto savivaldybės administracijos pastato renovacija, apimanti stogo remontą, fasado šiltinimą ir vidaus patalpų atnaujinimą.",
        procuring_organization=ProcuringOrganization(
            name="Vilniaus miesto savivaldybės administracija",
            code="188710061",
            email="info@vilnius.lt",
            phone="+370 5 211 2000",
        ),
        procurement_type="Atviras konkursas",
        estimated_value=EstimatedValue(
            amount=1_250_000.00,
            currency="EUR",
            vat_included=True,
            vat_amount=217_355.37,
        ),
        deadlines=Deadlines(
            submission_deadline="2026-03-15",
            questions_deadline="2026-03-01",
            contract_duration="24 mėnesiai",
            execution_deadline="2028-03-15",
        ),
        key_requirements=[
            "Pastato stogo renovacija pagal STR 2.05.01:2013",
            "Fasado šiltinimas ne mažiau kaip 150mm mineralinės vatos",
            "Vidaus patalpų atnaujinimas (1-3 aukštai)",
            "Priešgaisrinės saugos sistemų atnaujinimas",
        ],
        qualification_requirements=QualificationRequirements(
            financial=[
                "Metinė apyvarta ne mažesnė kaip 2,000,000 EUR",
                "Finansinių įsipareigojimų vykdymo garantija",
            ],
            technical=[
                "Turimų kvalifikuotų darbuotojų skaičius ne mažesnis kaip 20",
                "ISO 9001 kokybės vadybos sertifikatas",
            ],
            experience=[
                "Ne mažiau kaip 3 panašūs objektai per paskutinius 5 metus",
                "Bent vienas objektas, kurio vertė ne mažesnė kaip 500,000 EUR",
            ],
            other=[
                "Teisė verstis atitinkama veikla",
            ],
        ),
        evaluation_criteria=[
            EvaluationCriterion(
                criterion="Kaina",
                weight_percent=60.0,
                description="Mažiausia pasiūlyta kaina",
            ),
            EvaluationCriterion(
                criterion="Techniniai privalumai",
                weight_percent=25.0,
                description="Papildomi techniniai sprendimai",
            ),
            EvaluationCriterion(
                criterion="Terminai",
                weight_percent=15.0,
                description="Trumpesnis darbų atlikimo laikotarpis",
            ),
        ],
        restrictions_and_prohibitions=[
            "Draudžiama naudoti azbestą turinčias medžiagas",
            "Subrangovų dalis negali viršyti 40%",
        ],
        lot_structure=[
            LotInfo(
                lot_number=1,
                description="Stogo renovacija",
                estimated_value=450_000.00,
            ),
            LotInfo(
                lot_number=2,
                description="Fasado šiltinimas",
                estimated_value=550_000.00,
            ),
            LotInfo(
                lot_number=3,
                description="Vidaus patalpų atnaujinimas",
                estimated_value=250_000.00,
            ),
        ],
        special_conditions=[
            "Darbai vykdomi nepažeidžiant pastato veiklos",
            "Triukšmingi darbai tik darbo dienomis 8:00-18:00",
            "Privaloma draudimo polisas ne mažesniam kaip 500,000 EUR sumai",
        ],
        source_documents=[
            SourceDocument(filename="techninė_specifikacija.pdf", type="technical_spec", pages=45),
            SourceDocument(filename="sutarties_projektas.pdf", type="contract", pages=20),
        ],
        confidence_notes=[
            "PVM suma apskaičiuota pagal standartinį 21% tarifą",
            "Sutarties trukmė nurodyta preliminariai",
        ],
    )


@pytest.fixture
def full_qa() -> QAEvaluation:
    """QA evaluation with full data."""
    return QAEvaluation(
        completeness_score=0.85,
        missing_fields=["subrangovų reikalavimai"],
        conflicts=["Sutarties trukmė skiriasi tarp dokumentų (24 vs 18 mėnesių)"],
        suggestions=[
            "Patikrinti ar nurodytos kainos atitinka rinkos kainas",
            "Papildyti techninę specifikaciją energetinio naudingumo reikalavimais",
        ],
    )


@pytest.fixture
def minimal_report() -> AggregatedReport:
    """Report with minimal data (many None fields)."""
    return AggregatedReport(
        project_summary="Minimalus pirkimo aprašymas.",
    )


@pytest.fixture
def minimal_qa() -> QAEvaluation:
    """QA with minimal data."""
    return QAEvaluation(
        completeness_score=0.3,
        missing_fields=[
            "procuring_organization",
            "estimated_value",
            "deadlines",
            "evaluation_criteria",
        ],
    )


@pytest.fixture
def low_score_qa() -> QAEvaluation:
    """QA with low score (red)."""
    return QAEvaluation(
        completeness_score=0.2,
        missing_fields=["almost_everything"],
        conflicts=["Major conflict found"],
        suggestions=["Reanalyze documents"],
    )


@pytest.fixture
def high_score_qa() -> QAEvaluation:
    """QA with high score (green)."""
    return QAEvaluation(
        completeness_score=0.95,
    )


# ── Helper function tests ─────────────────────────────────────────────────────


class TestFormatValue:
    def test_none_value(self):
        assert _format_value(None) == "Nenurodyta"

    def test_none_amount(self):
        val = EstimatedValue(amount=None, currency="EUR")
        assert _format_value(val) == "Nenurodyta"

    def test_with_vat(self):
        val = EstimatedValue(amount=125_000.00, currency="EUR", vat_included=True)
        result = _format_value(val)
        assert "125,000.00" in result
        assert "EUR" in result
        assert "su PVM" in result

    def test_without_vat(self):
        val = EstimatedValue(amount=100_000.00, currency="EUR", vat_included=False)
        result = _format_value(val)
        assert "be PVM" in result

    def test_with_vat_amount(self):
        val = EstimatedValue(
            amount=125_000.00, currency="EUR", vat_included=True, vat_amount=21_694.21
        )
        result = _format_value(val)
        assert "PVM:" in result
        assert "21,694.21" in result


class TestFormatDate:
    def test_none(self):
        assert _format_date(None) == "Nenurodyta"

    def test_empty_string(self):
        assert _format_date("") == "Nenurodyta"

    def test_valid_date(self):
        result = _format_date("2026-03-15")
        assert "2026" in result
        assert "kovo" in result
        assert "15" in result

    def test_january(self):
        result = _format_date("2026-01-01")
        assert "sausio" in result

    def test_invalid_date(self):
        result = _format_date("not-a-date")
        assert result == "not-a-date"


class TestQAColor:
    def test_green(self):
        assert _qa_color(0.9) == "green"
        assert _qa_color(0.85) == "green"

    def test_orange(self):
        assert _qa_color(0.7) == "orange"
        assert _qa_color(0.6) == "orange"

    def test_red(self):
        assert _qa_color(0.3) == "red"
        assert _qa_color(0.0) == "red"

    def test_boundaries(self):
        assert _qa_color(0.8) == "orange"  # 0.8 is NOT > 0.8
        assert _qa_color(0.5) == "red"     # 0.5 is NOT > 0.5
        assert _qa_color(0.81) == "green"
        assert _qa_color(0.51) == "orange"


class TestSeverityColor:
    def test_info(self):
        r, g, b = _severity_color("info")
        assert b > r  # Blue dominant

    def test_warning(self):
        r, g, b = _severity_color("warning")
        assert r > b  # Warm color

    def test_conflict(self):
        r, g, b = _severity_color("conflict")
        assert r > g and r > b  # Red dominant

    def test_unknown(self):
        r, g, b = _severity_color("unknown")
        assert (r, g, b) == (0, 0, 0)  # Black fallback


# ── PDF Export tests ───────────────────────────────────────────────────────────


class TestExportPDF:
    @pytest.mark.asyncio
    async def test_pdf_full_data(self, full_report, full_qa):
        """PDF generation with full data produces a valid file."""
        path = await export_pdf(full_report, full_qa, model_used="claude-sonnet-4")
        try:
            assert path.exists()
            assert path.suffix == ".pdf"
            size = path.stat().st_size
            assert size > 0
            # Check PDF magic bytes
            with open(path, "rb") as f:
                header = f.read(5)
            assert header == b"%PDF-"
        finally:
            path.unlink(missing_ok=True)

    @pytest.mark.asyncio
    async def test_pdf_minimal_data(self, minimal_report, minimal_qa):
        """PDF generation with minimal data (many None fields)."""
        path = await export_pdf(minimal_report, minimal_qa)
        try:
            assert path.exists()
            assert path.suffix == ".pdf"
            size = path.stat().st_size
            assert size > 0
            with open(path, "rb") as f:
                header = f.read(5)
            assert header == b"%PDF-"
        finally:
            path.unlink(missing_ok=True)

    @pytest.mark.asyncio
    async def test_pdf_no_model(self, full_report, full_qa):
        """PDF without model_used still generates."""
        path = await export_pdf(full_report, full_qa, model_used="")
        try:
            assert path.exists()
            assert path.stat().st_size > 0
        finally:
            path.unlink(missing_ok=True)

    @pytest.mark.asyncio
    async def test_pdf_low_qa_score(self, full_report, low_score_qa):
        """PDF with low QA score still generates."""
        path = await export_pdf(full_report, low_score_qa, model_used="test")
        try:
            assert path.exists()
            assert path.stat().st_size > 0
        finally:
            path.unlink(missing_ok=True)

    @pytest.mark.asyncio
    async def test_pdf_high_qa_score(self, minimal_report, high_score_qa):
        """PDF with high QA score."""
        path = await export_pdf(minimal_report, high_score_qa)
        try:
            assert path.exists()
            assert path.stat().st_size > 0
        finally:
            path.unlink(missing_ok=True)


# ── DOCX Export tests ──────────────────────────────────────────────────────────


class TestExportDOCX:
    @pytest.mark.asyncio
    async def test_docx_full_data(self, full_report, full_qa):
        """DOCX generation with full data produces a valid file."""
        path = await export_docx(full_report, full_qa, model_used="claude-sonnet-4")
        try:
            assert path.exists()
            assert path.suffix == ".docx"
            size = path.stat().st_size
            assert size > 0
            # DOCX files are ZIP archives — check magic bytes
            with open(path, "rb") as f:
                header = f.read(4)
            assert header == b"PK\x03\x04"
        finally:
            path.unlink(missing_ok=True)

    @pytest.mark.asyncio
    async def test_docx_minimal_data(self, minimal_report, minimal_qa):
        """DOCX generation with minimal data (many None fields)."""
        path = await export_docx(minimal_report, minimal_qa)
        try:
            assert path.exists()
            assert path.suffix == ".docx"
            size = path.stat().st_size
            assert size > 0
            with open(path, "rb") as f:
                header = f.read(4)
            assert header == b"PK\x03\x04"
        finally:
            path.unlink(missing_ok=True)

    @pytest.mark.asyncio
    async def test_docx_no_model(self, full_report, full_qa):
        """DOCX without model_used still generates."""
        path = await export_docx(full_report, full_qa, model_used="")
        try:
            assert path.exists()
            assert path.stat().st_size > 0
        finally:
            path.unlink(missing_ok=True)

    @pytest.mark.asyncio
    async def test_docx_low_qa_score(self, full_report, low_score_qa):
        """DOCX with low QA score still generates."""
        path = await export_docx(full_report, low_score_qa, model_used="test")
        try:
            assert path.exists()
            assert path.stat().st_size > 0
        finally:
            path.unlink(missing_ok=True)

    @pytest.mark.asyncio
    async def test_docx_contains_content(self, full_report, full_qa):
        """DOCX contains expected content."""
        from docx import Document

        path = await export_docx(full_report, full_qa, model_used="test-model")
        try:
            doc = Document(str(path))
            full_text = "\n".join(p.text for p in doc.paragraphs)
            # Check key sections exist
            assert "Projekto santrauka" in full_text
            assert "Vilniaus miesto" in full_text
            assert "Vertinimo kriterijai" in full_text
            assert "Kokybės vertinimas" in full_text
            assert "85%" in full_text  # QA score
        finally:
            path.unlink(missing_ok=True)

    @pytest.mark.asyncio
    async def test_docx_tables_present(self, full_report, full_qa):
        """DOCX contains tables for criteria and lots."""
        from docx import Document

        path = await export_docx(full_report, full_qa)
        try:
            doc = Document(str(path))
            # Should have at least 2 tables (criteria + lots)
            assert len(doc.tables) >= 2
            # Check evaluation criteria table
            criteria_table = doc.tables[0]
            assert criteria_table.rows[0].cells[0].text == "Kriterijus"
            assert len(criteria_table.rows) == 4  # header + 3 criteria
        finally:
            path.unlink(missing_ok=True)
