# backend/app/services/exporter.py
# PDF and DOCX report export service
# Generates formatted Lithuanian procurement analysis reports
# Related: models/schemas.py (AggregatedReport, QAEvaluation)

import logging
from pathlib import Path
import tempfile
from datetime import datetime

from app.models.schemas import (
    AggregatedReport,
    ConfidenceNote,
    EstimatedValue,
    QAEvaluation,
)

logger = logging.getLogger(__name__)

# Lithuanian month names (genitive case for dates)
_LT_MONTHS_GENITIVE = {
    1: "sausio",
    2: "vasario",
    3: "kovo",
    4: "balandžio",
    5: "gegužės",
    6: "birželio",
    7: "liepos",
    8: "rugpjūčio",
    9: "rugsėjo",
    10: "spalio",
    11: "lapkričio",
    12: "gruodžio",
}

_NOT_SPECIFIED = "Nenurodyta"


# ── Helper functions ───────────────────────────────────────────────────────────


def _format_value(value: EstimatedValue | None) -> str:
    """Format estimated value: '125,000.00 EUR (su PVM)' or 'Nenurodyta'."""
    if value is None or value.amount is None:
        return _NOT_SPECIFIED
    formatted = f"{value.amount:,.2f} {value.currency}"
    if value.vat_included is True:
        formatted += " (su PVM)"
    elif value.vat_included is False:
        formatted += " (be PVM)"
    if value.vat_amount is not None:
        formatted += f", PVM: {value.vat_amount:,.2f} {value.currency}"
    return formatted


def _format_date(date_str: str | None) -> str:
    """Format ISO date to Lithuanian: '2026-03-15' → '2026 m. kovo 15 d.' or 'Nenurodyta'."""
    if not date_str:
        return _NOT_SPECIFIED
    try:
        dt = datetime.fromisoformat(date_str)
        month_name = _LT_MONTHS_GENITIVE.get(dt.month, str(dt.month))
        return f"{dt.year} m. {month_name} {dt.day} d."
    except (ValueError, TypeError):
        return date_str  # Return raw string if can't parse


def _qa_color(score: float) -> str:
    """Return color name for QA score."""
    if score > 0.8:
        return "green"
    elif score > 0.5:
        return "orange"
    else:
        return "red"


def _severity_color(severity: str) -> tuple:
    """Return RGB tuple for confidence note severity."""
    mapping = {
        "info": (0, 0, 180),       # Blue
        "warning": (200, 150, 0),   # Orange/yellow
        "conflict": (200, 0, 0),    # Red
    }
    return mapping.get(severity.lower(), (0, 0, 0))


def _risk_severity_color(severity: str) -> tuple:
    """Return RGB tuple for risk severity."""
    mapping = {
        "low": (46, 125, 50),        # Green
        "medium": (230, 81, 0),      # Orange
        "high": (198, 40, 40),       # Red
        "critical": (136, 14, 79),   # Dark magenta
    }
    return mapping.get(severity.lower(), (0, 0, 0))


def _or_na(value: str | None) -> str:
    """Return value or 'Nenurodyta' if None/empty."""
    return value if value else _NOT_SPECIFIED


def _format_org_contact(org) -> str:
    """Format organization contact info from expanded fields."""
    parts = []
    if org.contact_person:
        parts.append(org.contact_person)
    if org.phone:
        parts.append(f"Tel: {org.phone}")
    if org.email:
        parts.append(f"El. paštas: {org.email}")
    if org.website:
        parts.append(f"Svetainė: {org.website}")
    return "; ".join(parts) if parts else _NOT_SPECIFIED


def _format_org_address(org) -> str:
    """Format organization address from expanded fields."""
    parts = []
    if org.address:
        parts.append(org.address)
    if org.city:
        parts.append(org.city)
    if org.country:
        parts.append(org.country)
    return ", ".join(parts) if parts else _NOT_SPECIFIED


def _parse_confidence_notes(notes: list) -> list[ConfidenceNote]:
    """Parse confidence_notes which may be strings or ConfidenceNote objects."""
    result = []
    for note in notes:
        if isinstance(note, ConfidenceNote):
            result.append(note)
        elif isinstance(note, str):
            result.append(ConfidenceNote(note=note, severity="info"))
        elif isinstance(note, dict):
            result.append(ConfidenceNote(**note))
        else:
            result.append(ConfidenceNote(note=str(note), severity="info"))
    return result


# ── PDF Export ─────────────────────────────────────────────────────────────────


async def export_pdf(
    report: AggregatedReport,
    qa: QAEvaluation,
    model_used: str = "",
) -> Path:
    """
    Generate PDF report using reportlab.
    Returns path to generated PDF file.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib.colors import HexColor, Color
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
    )
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    # Create temp file
    tmp = tempfile.NamedTemporaryFile(
        suffix=".pdf", prefix="procurement_report_", delete=False
    )
    tmp.close()
    pdf_path = Path(tmp.name)

    logger.info("Generating PDF report: %s", pdf_path)

    doc = SimpleDocTemplate(
        str(pdf_path),
        pagesize=A4,
        topMargin=20 * mm,
        bottomMargin=25 * mm,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
    )

    styles = getSampleStyleSheet()

    # Custom styles
    styles.add(
        ParagraphStyle(
            "TitleLT",
            parent=styles["Title"],
            fontSize=18,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            "SubtitleLT",
            parent=styles["Normal"],
            fontSize=10,
            textColor=HexColor("#666666"),
            alignment=TA_CENTER,
            spaceAfter=12,
        )
    )
    styles.add(
        ParagraphStyle(
            "Heading2LT",
            parent=styles["Heading2"],
            fontSize=13,
            spaceBefore=12,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            "Heading3LT",
            parent=styles["Heading3"],
            fontSize=11,
            spaceBefore=8,
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            "BulletLT",
            parent=styles["Normal"],
            leftIndent=20,
            bulletIndent=10,
            spaceBefore=2,
            spaceAfter=2,
        )
    )
    styles.add(
        ParagraphStyle(
            "FooterLT",
            parent=styles["Normal"],
            fontSize=8,
            textColor=HexColor("#999999"),
            alignment=TA_CENTER,
        )
    )

    elements = []
    now_str = datetime.now().strftime("%Y-%m-%d %H:%M")
    section = 0

    def _heading(title: str):
        nonlocal section
        section += 1
        elements.append(Paragraph(f"{section}. {title}", styles["Heading2LT"]))

    def _subheading(title: str):
        elements.append(Paragraph(title, styles["Heading3LT"]))

    def _text(text: str):
        elements.append(Paragraph(text, styles["Normal"]))

    def _bold_text(label: str, value: str):
        elements.append(Paragraph(f"<b>{label}:</b> {value}", styles["Normal"]))

    def _bullet(text: str):
        elements.append(Paragraph(f"• {text}", styles["BulletLT"]))

    def _spacer():
        elements.append(Spacer(1, 3 * mm))

    # ── Title
    elements.append(Paragraph("Viešojo pirkimo analizė", styles["TitleLT"]))
    subtitle = f"Sugeneruota: {now_str}"
    if model_used:
        subtitle += f" | Modelis: {model_used}"
    elements.append(Paragraph(subtitle, styles["SubtitleLT"]))
    elements.append(Spacer(1, 6 * mm))

    # ── 1. Pagrindinė informacija
    _heading("Pagrindinė informacija")
    if report.project_title:
        _bold_text("Projekto pavadinimas", report.project_title)
    org = report.procuring_organization
    if org and org.name:
        _bold_text("Perkančioji organizacija", org.name)
    _bold_text("Projekto vertė", _format_value(report.estimated_value))
    dl = report.deadlines
    if dl and dl.submission_deadline:
        _bold_text("Dokumentų pateikimo terminas", _format_date(dl.submission_deadline))
    if report.procurement_reference:
        _bold_text("CVP kodas", report.procurement_reference)
    if report.cpv_codes:
        _bold_text("CPV kodai", "; ".join(report.cpv_codes))
    _spacer()

    # ── 2. Projekto santrauka
    _heading("Projekto santrauka")
    _text(_or_na(report.project_summary))
    if report.nuts_codes:
        _bold_text("NUTS kodai", "; ".join(report.nuts_codes))
    if report.procurement_law:
        _bold_text("Teisės aktas", report.procurement_law)
    _spacer()

    # ── 3. Perkančioji organizacija
    _heading("Perkančioji organizacija")
    if org:
        _bold_text("Pavadinimas", _or_na(org.name))
        _bold_text("Kodas", _or_na(org.code))
        if org.organization_type:
            _bold_text("Tipas", org.organization_type)
        _bold_text("Adresas", _format_org_address(org))
        _bold_text("Kontaktai", _format_org_contact(org))
    else:
        _text(_NOT_SPECIFIED)
    _spacer()

    # ── 4. Pirkimo būdas
    if report.procurement_type:
        _heading("Pirkimo būdas")
        _text(report.procurement_type)
        _spacer()

    # ── 5. Finansinės sąlygos
    ft = report.financial_terms
    if ft:
        _heading("Finansinės sąlygos")
        if ft.payment_terms:
            _bold_text("Mokėjimo sąlygos", ft.payment_terms)
        if ft.advance_payment:
            _bold_text("Avansinis mokėjimas", ft.advance_payment)
        if ft.guarantee_requirements:
            _bold_text("Garantijos reikalavimai", ft.guarantee_requirements)
        if ft.guarantee_amount:
            _bold_text("Garantijos dydis", ft.guarantee_amount)
        if ft.price_adjustment:
            _bold_text("Kainos keitimo sąlygos", ft.price_adjustment)
        if ft.insurance_requirements:
            _bold_text("Draudimo reikalavimai", ft.insurance_requirements)
        if ft.penalty_clauses:
            _subheading("Baudos ir netesybos:")
            for p in ft.penalty_clauses:
                _bullet(p)
        _spacer()

    # ── 6. Terminai
    _heading("Terminai")
    if dl:
        _bold_text("Pasiūlymų pateikimas", _format_date(dl.submission_deadline))
        _bold_text("Klausimų pateikimas", _format_date(dl.questions_deadline))
        _bold_text("Sutarties trukmė", _or_na(dl.contract_duration))
        _bold_text("Darbų atlikimas", _format_date(dl.execution_deadline))
        if dl.offer_validity:
            _bold_text("Pasiūlymo galiojimas", dl.offer_validity)
        if dl.contract_start:
            _bold_text("Sutarties pradžia", dl.contract_start)
        if dl.extension_options:
            _bold_text("Pratęsimo galimybės", dl.extension_options)
    else:
        _text(_NOT_SPECIFIED)
    _spacer()

    # ── 7. Techninė specifikacija
    _heading("Techninė specifikacija")
    if report.technical_specifications:
        for ts in report.technical_specifications:
            mandatory_tag = "PRIVALOMA" if ts.mandatory else "PAGEIDAUJAMA"
            _bullet(f"[{mandatory_tag}] {ts.description}")
            if ts.details:
                elements.append(
                    Paragraph(f"    ↳ {ts.details}", styles["BulletLT"])
                )
    elif report.key_requirements:
        for req in report.key_requirements:
            _bullet(req)
    else:
        _text(_NOT_SPECIFIED)
    _spacer()

    # ── 8. Pagrindiniai reikalavimai (if both exist)
    if report.technical_specifications and report.key_requirements:
        _heading("Kiti pagrindiniai reikalavimai")
        for req in report.key_requirements:
            _bullet(req)
        _spacer()

    # ── 9. Kvalifikacijos reikalavimai
    _heading("Kvalifikacijos reikalavimai")
    qr = report.qualification_requirements
    if qr:
        for group_name, group_label in [
            ("financial", "Finansiniai"),
            ("technical", "Techniniai"),
            ("experience", "Patirties"),
            ("personnel", "Personalo"),
            ("exclusion_grounds", "Pašalinimo pagrindai"),
            ("required_documents", "Reikalaujami dokumentai"),
            ("other", "Kiti"),
        ]:
            items = getattr(qr, group_name, [])
            if items:
                _subheading(f"{group_label}:")
                for item in items:
                    _bullet(item)
    else:
        _text(_NOT_SPECIFIED)
    _spacer()

    # ── 10. Vertinimo kriterijai (TABLE)
    _heading("Vertinimo kriterijai")
    if report.evaluation_criteria:
        table_data = [["Kriterijus", "Svoris (%)", "Aprašymas"]]
        for ec in report.evaluation_criteria:
            weight = f"{ec.weight_percent:.1f}" if ec.weight_percent is not None else "-"
            table_data.append([ec.criterion, weight, _or_na(ec.description)])

        col_widths = [150, 60, 260]
        table = Table(table_data, colWidths=col_widths)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), HexColor("#4472C4")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTSIZE", (0, 0), (-1, 0), 9),
                    ("FONTSIZE", (0, 1), (-1, -1), 8),
                    ("ALIGN", (1, 0), (1, -1), "CENTER"),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, HexColor("#F2F2F2")]),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ]
            )
        )
        elements.append(table)
    else:
        _text(_NOT_SPECIFIED)
    _spacer()

    # ── 11. Pasiūlymo pateikimas
    sr = report.submission_requirements
    if sr:
        _heading("Pasiūlymo pateikimas")
        if sr.submission_method:
            _bold_text("Pateikimo būdas", sr.submission_method)
        if sr.submission_language:
            _bold_text("Kalbos", ", ".join(sr.submission_language))
        if sr.required_format:
            _bold_text("Formatas", sr.required_format)
        if sr.envelope_system:
            _bold_text("Vokelių sistema", sr.envelope_system)
        if sr.variants_allowed is not None:
            _bold_text("Alternatyvūs pasiūlymai", "Leidžiami" if sr.variants_allowed else "Neleidžiami")
        if sr.joint_bidding:
            _bold_text("Jungtiniai pasiūlymai", sr.joint_bidding)
        if sr.subcontracting:
            _bold_text("Subrangos sąlygos", sr.subcontracting)
        _spacer()

    # ── 12. Rizikos tiekėjui (TABLE)
    if report.risk_factors:
        _heading("Rizikos tiekėjui")
        table_data = [["Rizika", "Lygis", "Rekomendacija"]]
        for rf in report.risk_factors:
            table_data.append([rf.risk, rf.severity.upper(), _or_na(rf.recommendation)])

        col_widths = [180, 55, 235]
        table = Table(table_data, colWidths=col_widths)

        # Color-code severity in the table
        table_style_cmds = [
            ("BACKGROUND", (0, 0), (-1, 0), HexColor("#C62828")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, 0), 9),
            ("FONTSIZE", (0, 1), (-1, -1), 8),
            ("ALIGN", (1, 0), (1, -1), "CENTER"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, HexColor("#FFF3F3")]),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ]
        table.setStyle(TableStyle(table_style_cmds))
        elements.append(table)
        _spacer()

    # ── 13. Lotai (TABLE)
    if report.lot_structure:
        _heading("Lotai")
        lot_data = [["Nr.", "Aprašymas", "Vertė (EUR)"]]
        for lot in report.lot_structure:
            val = f"{lot.estimated_value:,.2f}" if lot.estimated_value is not None else "-"
            lot_data.append([str(lot.lot_number), lot.description, val])

        lot_table = Table(lot_data, colWidths=[40, 320, 100])
        lot_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), HexColor("#4472C4")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTSIZE", (0, 0), (-1, 0), 9),
                    ("FONTSIZE", (0, 1), (-1, -1), 8),
                    ("ALIGN", (0, 0), (0, -1), "CENTER"),
                    ("ALIGN", (2, 0), (2, -1), "RIGHT"),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, HexColor("#F2F2F2")]),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )
        elements.append(lot_table)
        _spacer()

    # ── 14. Specialios sąlygos
    _heading("Specialios sąlygos")
    if report.special_conditions:
        for cond in report.special_conditions:
            _bullet(cond)
    else:
        _text(_NOT_SPECIFIED)
    _spacer()

    # ── 15. Apribojimai ir draudimai
    _heading("Apribojimai ir draudimai")
    if report.restrictions_and_prohibitions:
        for r in report.restrictions_and_prohibitions:
            _bullet(r)
    else:
        _text(_NOT_SPECIFIED)
    _spacer()

    # ── 16. Apeliavimas
    if report.appeal_procedures:
        _heading("Apeliavimo procedūra")
        _text(report.appeal_procedures)
        _spacer()

    # ── 17. Pastabos ir patikimumas
    _heading("Pastabos ir patikimumas")
    if report.confidence_notes:
        notes = _parse_confidence_notes(report.confidence_notes)
        for cn in notes:
            r, g, b = _severity_color(cn.severity)
            color_hex = f"#{r:02x}{g:02x}{b:02x}"
            elements.append(
                Paragraph(
                    f'<font color="{color_hex}">[{cn.severity.upper()}]</font> {cn.note}',
                    styles["Normal"],
                )
            )
    else:
        _text("Pastabų nėra")
    _spacer()

    # ── 18. Kokybės vertinimas (QA)
    _heading("Kokybės vertinimas")
    qa_color = _qa_color(qa.completeness_score)
    color_map = {"green": "#2E7D32", "orange": "#E65100", "red": "#C62828"}
    qa_hex = color_map.get(qa_color, "#000000")

    elements.append(
        Paragraph(
            f'<b>Užbaigtumo balas:</b> <font color="{qa_hex}">'
            f"{qa.completeness_score:.0%}</font>",
            styles["Normal"],
        )
    )

    if qa.missing_fields:
        _subheading("Trūkstami laukai:")
        for mf in qa.missing_fields:
            _bullet(mf)

    if qa.conflicts:
        _subheading("Prieštaravimai:")
        for conflict in qa.conflicts:
            elements.append(
                Paragraph(
                    f'<font color="#C62828">• {conflict}</font>',
                    styles["BulletLT"],
                )
            )

    if qa.suggestions:
        _subheading("Pasiūlymai:")
        for sug in qa.suggestions:
            _bullet(sug)

    elements.append(Spacer(1, 6 * mm))

    # ── Footer
    footer_text = f"Sugeneruota: {now_str}"
    if model_used:
        footer_text += f" | Modelis: {model_used}"
    elements.append(Paragraph(footer_text, styles["FooterLT"]))

    # Build
    doc.build(elements)
    logger.info("PDF report generated: %s (%d bytes)", pdf_path, pdf_path.stat().st_size)
    return pdf_path


# ── DOCX Export ────────────────────────────────────────────────────────────────


async def export_docx(
    report: AggregatedReport,
    qa: QAEvaluation,
    model_used: str = "",
) -> Path:
    """
    Generate DOCX report using python-docx.
    Returns path to generated DOCX file.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    tmp = tempfile.NamedTemporaryFile(
        suffix=".docx", prefix="procurement_report_", delete=False
    )
    tmp.close()
    docx_path = Path(tmp.name)

    logger.info("Generating DOCX report: %s", docx_path)

    doc = Document()
    now_str = datetime.now().strftime("%Y-%m-%d %H:%M")
    section_num = 0

    def _heading(title: str):
        nonlocal section_num
        section_num += 1
        doc.add_heading(f"{section_num}. {title}", level=1)

    def _subheading(title: str):
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(11)

    def _bold_para(label: str, value: str):
        p = doc.add_paragraph()
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(value)

    # ── Title
    title_para = doc.add_heading("Viešojo pirkimo analizė", level=0)
    subtitle = f"Sugeneruota: {now_str}"
    if model_used:
        subtitle += f" | Modelis: {model_used}"
    sub_para = doc.add_paragraph(subtitle)
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub_para.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(102, 102, 102)

    # ── 1. Pagrindinė informacija
    _heading("Pagrindinė informacija")
    org = report.procuring_organization
    dl = report.deadlines
    if report.project_title:
        _bold_para("Projekto pavadinimas", report.project_title)
    if org and org.name:
        _bold_para("Perkančioji organizacija", org.name)
    _bold_para("Projekto vertė", _format_value(report.estimated_value))
    if dl and dl.submission_deadline:
        _bold_para("Dokumentų pateikimo terminas", _format_date(dl.submission_deadline))
    if report.procurement_reference:
        _bold_para("CVP kodas", report.procurement_reference)
    if report.cpv_codes:
        _bold_para("CPV kodai", "; ".join(report.cpv_codes))

    # ── 2. Projekto santrauka
    _heading("Projekto santrauka")
    doc.add_paragraph(_or_na(report.project_summary))
    if report.nuts_codes:
        _bold_para("NUTS kodai", "; ".join(report.nuts_codes))
    if report.procurement_law:
        _bold_para("Teisės aktas", report.procurement_law)

    # ── 3. Perkančioji organizacija
    _heading("Perkančioji organizacija")
    if org:
        _bold_para("Pavadinimas", _or_na(org.name))
        _bold_para("Kodas", _or_na(org.code))
        if org.organization_type:
            _bold_para("Tipas", org.organization_type)
        _bold_para("Adresas", _format_org_address(org))
        _bold_para("Kontaktai", _format_org_contact(org))
    else:
        doc.add_paragraph(_NOT_SPECIFIED)

    # ── 4. Pirkimo būdas
    if report.procurement_type:
        _heading("Pirkimo būdas")
        doc.add_paragraph(report.procurement_type)

    # ── 5. Finansinės sąlygos
    ft = report.financial_terms
    if ft:
        _heading("Finansinės sąlygos")
        if ft.payment_terms:
            _bold_para("Mokėjimo sąlygos", ft.payment_terms)
        if ft.advance_payment:
            _bold_para("Avansinis mokėjimas", ft.advance_payment)
        if ft.guarantee_requirements:
            _bold_para("Garantijos reikalavimai", ft.guarantee_requirements)
        if ft.guarantee_amount:
            _bold_para("Garantijos dydis", ft.guarantee_amount)
        if ft.price_adjustment:
            _bold_para("Kainos keitimo sąlygos", ft.price_adjustment)
        if ft.insurance_requirements:
            _bold_para("Draudimo reikalavimai", ft.insurance_requirements)
        if ft.penalty_clauses:
            _subheading("Baudos ir netesybos:")
            for p in ft.penalty_clauses:
                doc.add_paragraph(p, style="List Bullet")

    # ── 6. Terminai
    _heading("Terminai")
    if dl:
        _bold_para("Pasiūlymų pateikimas", _format_date(dl.submission_deadline))
        _bold_para("Klausimų pateikimas", _format_date(dl.questions_deadline))
        _bold_para("Sutarties trukmė", _or_na(dl.contract_duration))
        _bold_para("Darbų atlikimas", _format_date(dl.execution_deadline))
        if dl.offer_validity:
            _bold_para("Pasiūlymo galiojimas", dl.offer_validity)
        if dl.contract_start:
            _bold_para("Sutarties pradžia", dl.contract_start)
        if dl.extension_options:
            _bold_para("Pratęsimo galimybės", dl.extension_options)
    else:
        doc.add_paragraph(_NOT_SPECIFIED)

    # ── 7. Techninė specifikacija
    _heading("Techninė specifikacija")
    if report.technical_specifications:
        for ts in report.technical_specifications:
            tag = "PRIVALOMA" if ts.mandatory else "PAGEIDAUJAMA"
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"[{tag}] ")
            run.bold = True
            p.add_run(ts.description)
            if ts.details:
                detail_p = doc.add_paragraph(f"    ↳ {ts.details}")
                detail_p.paragraph_format.left_indent = Pt(36)
    elif report.key_requirements:
        for req in report.key_requirements:
            doc.add_paragraph(req, style="List Bullet")
    else:
        doc.add_paragraph(_NOT_SPECIFIED)

    # ── 8. Kiti pagrindiniai reikalavimai (if both exist)
    if report.technical_specifications and report.key_requirements:
        _heading("Kiti pagrindiniai reikalavimai")
        for req in report.key_requirements:
            doc.add_paragraph(req, style="List Bullet")

    # ── 9. Kvalifikacijos reikalavimai
    _heading("Kvalifikacijos reikalavimai")
    qr = report.qualification_requirements
    if qr:
        for group_name, group_label in [
            ("financial", "Finansiniai"),
            ("technical", "Techniniai"),
            ("experience", "Patirties"),
            ("personnel", "Personalo"),
            ("exclusion_grounds", "Pašalinimo pagrindai"),
            ("required_documents", "Reikalaujami dokumentai"),
            ("other", "Kiti"),
        ]:
            items = getattr(qr, group_name, [])
            if items:
                _subheading(f"{group_label}:")
                for item in items:
                    doc.add_paragraph(item, style="List Bullet")
    else:
        doc.add_paragraph(_NOT_SPECIFIED)

    # ── 10. Vertinimo kriterijai (TABLE)
    _heading("Vertinimo kriterijai")
    if report.evaluation_criteria:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Kriterijus"
        hdr[1].text = "Svoris (%)"
        hdr[2].text = "Aprašymas"

        for ec in report.evaluation_criteria:
            row = table.add_row().cells
            row[0].text = ec.criterion
            row[1].text = (
                f"{ec.weight_percent:.1f}" if ec.weight_percent is not None else "-"
            )
            row[2].text = _or_na(ec.description)
    else:
        doc.add_paragraph(_NOT_SPECIFIED)

    # ── 11. Pasiūlymo pateikimas
    sr = report.submission_requirements
    if sr:
        _heading("Pasiūlymo pateikimas")
        if sr.submission_method:
            _bold_para("Pateikimo būdas", sr.submission_method)
        if sr.submission_language:
            _bold_para("Kalbos", ", ".join(sr.submission_language))
        if sr.required_format:
            _bold_para("Formatas", sr.required_format)
        if sr.envelope_system:
            _bold_para("Vokelių sistema", sr.envelope_system)
        if sr.variants_allowed is not None:
            _bold_para("Alternatyvūs pasiūlymai", "Leidžiami" if sr.variants_allowed else "Neleidžiami")
        if sr.joint_bidding:
            _bold_para("Jungtiniai pasiūlymai", sr.joint_bidding)
        if sr.subcontracting:
            _bold_para("Subrangos sąlygos", sr.subcontracting)

    # ── 12. Rizikos tiekėjui (TABLE)
    if report.risk_factors:
        _heading("Rizikos tiekėjui")
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 2"
        hdr = table.rows[0].cells
        hdr[0].text = "Rizika"
        hdr[1].text = "Lygis"
        hdr[2].text = "Rekomendacija"

        for rf in report.risk_factors:
            row = table.add_row().cells
            row[0].text = rf.risk
            row[1].text = rf.severity.upper()
            row[2].text = _or_na(rf.recommendation)

    # ── 13. Lotai (TABLE)
    if report.lot_structure:
        _heading("Lotai")
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Nr."
        hdr[1].text = "Aprašymas"
        hdr[2].text = "Vertė (EUR)"

        for lot in report.lot_structure:
            row = table.add_row().cells
            row[0].text = str(lot.lot_number)
            row[1].text = lot.description
            row[2].text = (
                f"{lot.estimated_value:,.2f}" if lot.estimated_value is not None else "-"
            )

    # ── 14. Specialios sąlygos
    _heading("Specialios sąlygos")
    if report.special_conditions:
        for cond in report.special_conditions:
            doc.add_paragraph(cond, style="List Bullet")
    else:
        doc.add_paragraph(_NOT_SPECIFIED)

    # ── 15. Apribojimai ir draudimai
    _heading("Apribojimai ir draudimai")
    if report.restrictions_and_prohibitions:
        for r in report.restrictions_and_prohibitions:
            doc.add_paragraph(r, style="List Bullet")
    else:
        doc.add_paragraph(_NOT_SPECIFIED)

    # ── 16. Apeliavimas
    if report.appeal_procedures:
        _heading("Apeliavimo procedūra")
        doc.add_paragraph(report.appeal_procedures)

    # ── 17. Pastabos ir patikimumas
    _heading("Pastabos ir patikimumas")
    if report.confidence_notes:
        notes = _parse_confidence_notes(report.confidence_notes)
        for cn in notes:
            p = doc.add_paragraph()
            severity_run = p.add_run(f"[{cn.severity.upper()}] ")
            r, g, b = _severity_color(cn.severity)
            severity_run.font.color.rgb = RGBColor(r, g, b)
            severity_run.bold = True
            p.add_run(cn.note)
    else:
        doc.add_paragraph("Pastabų nėra")

    # ── 18. Kokybės vertinimas
    _heading("Kokybės vertinimas")

    # QA Score with color
    qa_p = doc.add_paragraph()
    qa_p.add_run("Užbaigtumo balas: ").bold = True
    score_run = qa_p.add_run(f"{qa.completeness_score:.0%}")
    score_run.bold = True
    qa_color = _qa_color(qa.completeness_score)
    color_map = {
        "green": RGBColor(0x2E, 0x7D, 0x32),
        "orange": RGBColor(0xE6, 0x51, 0x00),
        "red": RGBColor(0xC6, 0x28, 0x28),
    }
    score_run.font.color.rgb = color_map.get(qa_color, RGBColor(0, 0, 0))

    if qa.missing_fields:
        _subheading("Trūkstami laukai:")
        for mf in qa.missing_fields:
            doc.add_paragraph(mf, style="List Bullet")

    if qa.conflicts:
        _subheading("Prieštaravimai:")
        for conflict in qa.conflicts:
            cp = doc.add_paragraph(style="List Bullet")
            run = cp.add_run(conflict)
            run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)

    if qa.suggestions:
        _subheading("Pasiūlymai:")
        for sug in qa.suggestions:
            doc.add_paragraph(sug, style="List Bullet")

    # ── Footer
    doc.add_paragraph()  # spacer
    footer_text = f"Sugeneruota: {now_str}"
    if model_used:
        footer_text += f" | Modelis: {model_used}"
    footer_para = doc.add_paragraph(footer_text)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_para.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(153, 153, 153)

    doc.save(str(docx_path))
    logger.info("DOCX report generated: %s (%d bytes)", docx_path, docx_path.stat().st_size)
    return docx_path
