# backend/app/services/parser.py
# Document parsing service — fast parsers (PyMuPDF, python-docx) with Docling fallback
# Converts PDF, DOCX, XLSX, PPTX, images to markdown text
# Related: models/schemas.py, services/zip_extractor.py

import asyncio
import logging
import re
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Callable, Optional

from app.models.schemas import DocumentType

logger = logging.getLogger(__name__)


# Suppress known Docling regression: ListGroup warnings from msword_backend
# See: https://github.com/DS4SD/docling/issues/2967 (caused by PR #2665)
class _DoclingListWarningFilter(logging.Filter):
    def filter(self, record: logging.LogRecord) -> bool:
        msg = record.getMessage()
        if "list item" in msg.lower() and "ListGroup" in msg:
            return False
        if "List item not matching any insert condition" in msg:
            return False
        return True


logging.getLogger("docling.backend.msword_backend").addFilter(_DoclingListWarningFilter())


@dataclass
class ParsedDocument:
    filename: str
    content: str  # markdown text
    page_count: int
    file_size_bytes: int
    doc_type: DocumentType
    token_estimate: int  # len(content) // 4 rough estimate
    file_path: Optional[Path] = None  # original file path for multimodal OCR
    is_scanned: bool = False  # True = empty text, needs vision/OCR extraction


# ── Fast parsers (pypdf for PDF, python-docx for DOCX) ───────────────────────


def _parse_pdf_fast(file_path: Path) -> tuple[str, int]:
    """Parse PDF using pypdf — returns (markdown_text, page_count).

    Pure Python, no native DLLs — avoids Windows Application Control blocks.
    Fast for text-based PDFs.
    """
    from pypdf import PdfReader

    reader = PdfReader(str(file_path))
    page_count = len(reader.pages)

    parts: list[str] = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text.strip())

    markdown_text = "\n\n".join(parts)
    return markdown_text, page_count


def _parse_docx_fast(file_path: Path) -> tuple[str, int]:
    """Parse DOCX using python-docx — returns (markdown_text, page_count).

    Extracts paragraphs, tables, and basic formatting as markdown.
    """
    from docx import Document

    doc = Document(str(file_path))
    parts: list[str] = []

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            # Paragraph
            para = None
            for p in doc.paragraphs:
                if p._element is element:
                    para = p
                    break
            if para is None:
                continue

            text = para.text.strip()
            if not text:
                continue

            style_name = (para.style.name or "").lower() if para.style else ""

            if "heading 1" in style_name:
                parts.append(f"# {text}")
            elif "heading 2" in style_name:
                parts.append(f"## {text}")
            elif "heading 3" in style_name:
                parts.append(f"### {text}")
            elif "heading" in style_name:
                parts.append(f"#### {text}")
            elif "list" in style_name or "bullet" in style_name:
                parts.append(f"- {text}")
            else:
                parts.append(text)

        elif tag == "tbl":
            # Table
            for table in doc.tables:
                if table._element is element:
                    _render_table(table, parts)
                    break

    markdown_text = "\n\n".join(parts)

    # Estimate pages from content length (~3000 chars per page)
    page_count = max(len(markdown_text) // 3000, 1)

    return markdown_text, page_count


def _render_table(table, parts: list[str]) -> None:
    """Render a python-docx table as markdown."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(cells)

    if not rows:
        return

    # Header row
    parts.append("| " + " | ".join(rows[0]) + " |")
    parts.append("| " + " | ".join("---" for _ in rows[0]) + " |")

    # Data rows
    for row in rows[1:]:
        # Pad row to match header column count
        while len(row) < len(rows[0]):
            row.append("")
        parts.append("| " + " | ".join(row[: len(rows[0])]) + " |")


# ── Docling fallback (for images, PPTX, and complex formats) ─────────────────

_converter = None


def _get_converter():
    """Lazily initialize and cache the Docling DocumentConverter.

    Only used as fallback for formats not handled by fast parsers
    (images, PPTX, or when fast parsing fails).
    """
    global _converter
    if _converter is None:
        import os

        from docling.datamodel.base_models import InputFormat
        from docling.datamodel.pipeline_options import (
            AcceleratorOptions,
            PdfPipelineOptions,
            TableFormerMode,
            TableStructureOptions,
        )
        from docling.document_converter import (
            DocumentConverter,
            ImageFormatOption,
            PdfFormatOption,
        )

        from app.config import get_settings

        settings = get_settings()
        cpu_threads = os.cpu_count() or 4

        table_opts = TableStructureOptions(mode=TableFormerMode.FAST)
        accel_opts = AcceleratorOptions(num_threads=cpu_threads, device="cpu")

        pdf_opts = PdfPipelineOptions(
            do_ocr=False,
            do_table_structure=True,
            table_structure_options=table_opts,
            document_timeout=settings.parser_doc_timeout,
            force_backend_text=settings.parser_force_backend_text,
            accelerator_options=accel_opts,
        )

        _converter = DocumentConverter(
            format_options={
                InputFormat.PDF: PdfFormatOption(pipeline_options=pdf_opts),
                InputFormat.IMAGE: ImageFormatOption(pipeline_options=pdf_opts),
            }
        )

        logger.info(
            "Docling converter initialized (fallback): table_mode=fast, "
            "timeout=%ds, threads=%d",
            settings.parser_doc_timeout,
            cpu_threads,
        )
    return _converter


_ocr_converter = None


def _get_ocr_converter():
    """Lazily initialize Docling converter with OCR enabled (RapidOCR).

    Used ONLY as fallback for scanned files > 5MB that can't be sent via
    OpenRouter multimodal API.
    """
    global _ocr_converter
    if _ocr_converter is None:
        import os

        from docling.datamodel.base_models import InputFormat
        from docling.datamodel.pipeline_options import (
            AcceleratorOptions,
            PdfPipelineOptions,
            RapidOcrOptions,
            TableFormerMode,
            TableStructureOptions,
        )
        from docling.document_converter import (
            DocumentConverter,
            ImageFormatOption,
            PdfFormatOption,
        )

        from app.config import get_settings

        settings = get_settings()
        cpu_threads = os.cpu_count() or 4

        table_opts = TableStructureOptions(mode=TableFormerMode.FAST)
        accel_opts = AcceleratorOptions(num_threads=cpu_threads, device="cpu")
        ocr_opts = RapidOcrOptions()

        pdf_opts = PdfPipelineOptions(
            do_ocr=True,
            ocr_options=ocr_opts,
            do_table_structure=True,
            table_structure_options=table_opts,
            document_timeout=settings.parser_doc_timeout,
            accelerator_options=accel_opts,
        )

        _ocr_converter = DocumentConverter(
            format_options={
                InputFormat.PDF: PdfFormatOption(pipeline_options=pdf_opts),
                InputFormat.IMAGE: ImageFormatOption(pipeline_options=pdf_opts),
            }
        )

        logger.info(
            "Docling OCR converter initialized: ocr=RapidOCR, "
            "timeout=%ds, threads=%d",
            settings.parser_doc_timeout,
            cpu_threads,
        )
    return _ocr_converter


def parse_with_ocr(file_path: Path) -> tuple[str, int]:
    """Parse a scanned document using Docling with OCR enabled.

    Used as fallback for files > 5MB. Returns (markdown_text, page_count).
    """
    converter = _get_ocr_converter()
    result = converter.convert(str(file_path))

    from docling.datamodel.base_models import ConversionStatus

    if result.status == ConversionStatus.FAILURE:
        error_msgs = "; ".join(str(e) for e in (result.errors or []))
        raise RuntimeError(f"Docling OCR conversion failed: {error_msgs or 'unknown error'}")

    markdown_text = result.document.export_to_markdown()
    file_ext = file_path.suffix.lower()

    try:
        page_count = result.document.num_pages()
        if page_count == 0:
            page_count = _estimate_pages(markdown_text, file_ext)
    except Exception:
        page_count = _estimate_pages(markdown_text, file_ext)

    logger.info(
        "OCR parsing complete for %s: %d pages, %d chars",
        file_path.name, page_count, len(markdown_text),
    )
    return markdown_text, page_count


def _parse_with_docling(file_path: Path, file_ext: str) -> tuple[str, int]:
    """Parse using Docling — fallback for images, PPTX, and complex formats."""
    converter = _get_converter()
    result = converter.convert(str(file_path))

    from docling.datamodel.base_models import ConversionStatus

    if result.status == ConversionStatus.FAILURE:
        error_msgs = "; ".join(str(e) for e in (result.errors or []))
        raise RuntimeError(f"Docling conversion failed: {error_msgs or 'unknown error'}")

    markdown_text = result.document.export_to_markdown()

    try:
        page_count = result.document.num_pages()
        if page_count == 0:
            page_count = _estimate_pages(markdown_text, file_ext)
    except Exception:
        page_count = _estimate_pages(markdown_text, file_ext)

    return markdown_text, page_count


# ── Main parse function ──────────────────────────────────────────────────────


# Extensions handled by fast parsers (no Docling needed)
_FAST_PDF_EXTS = {".pdf"}
_FAST_DOCX_EXTS = {".docx"}
# Extensions that need Docling (images, PPTX, XLSX)
_DOCLING_EXTS = {".pptx", ".png", ".tiff", ".jpg", ".jpeg", ".xlsx"}
# Image extensions — always treated as scanned (need vision/OCR)
_IMAGE_EXTS = {".png", ".tiff", ".jpg", ".jpeg"}


async def parse_document(file_path: Path, filename: str) -> ParsedDocument:
    """Parse a single document — uses fast parser when possible, Docling as fallback.

    Strategy:
    - PDF → pypdf (fast, pure Python, ~0.1-0.5s)
    - DOCX → python-docx (fast, ~0.1s)
    - XLSX/PPTX/images → Docling (slow but necessary)
    - If fast parser fails → automatic Docling fallback
    """
    try:
        file_size = file_path.stat().st_size
    except (FileNotFoundError, OSError) as e:
        logger.warning("File not found or inaccessible: %s — %s", filename, e)
        error_content = f"[ERROR] File not found: {filename}"
        doc_type = classify_document(filename, "")
        return ParsedDocument(
            filename=filename,
            content=error_content,
            page_count=0,
            file_size_bytes=0,
            doc_type=doc_type,
            token_estimate=len(error_content) // 4,
            file_path=file_path,
        )

    file_ext = file_path.suffix.lower()
    start = time.perf_counter()

    try:
        loop = asyncio.get_running_loop()

        if file_ext in _FAST_PDF_EXTS:
            # Fast path: pypdf
            try:
                markdown_text, page_count = await loop.run_in_executor(
                    None, _parse_pdf_fast, file_path
                )
                parser_used = "pypdf"
            except Exception as e:
                logger.warning(
                    "pypdf failed for %s (%s), falling back to Docling", filename, e
                )
                markdown_text, page_count = await loop.run_in_executor(
                    None, _parse_with_docling, file_path, file_ext
                )
                parser_used = "docling-fallback"

        elif file_ext in _FAST_DOCX_EXTS:
            # Fast path: python-docx
            try:
                markdown_text, page_count = await loop.run_in_executor(
                    None, _parse_docx_fast, file_path
                )
                parser_used = "python-docx"
            except Exception as e:
                logger.warning(
                    "python-docx failed for %s (%s), falling back to Docling",
                    filename,
                    e,
                )
                markdown_text, page_count = await loop.run_in_executor(
                    None, _parse_with_docling, file_path, file_ext
                )
                parser_used = "docling-fallback"

        else:
            # Docling for everything else (images, PPTX, XLSX)
            markdown_text, page_count = await loop.run_in_executor(
                None, _parse_with_docling, file_path, file_ext
            )
            parser_used = "docling"

        elapsed = time.perf_counter() - start

        # Classify document type
        content_preview = markdown_text[:2000]
        doc_type = classify_document(filename, content_preview)

        # Token estimate (~4 chars per token)
        token_estimate = len(markdown_text) // 4

        # Detect scanned documents (empty/near-empty text)
        from app.config import get_settings
        settings = get_settings()
        is_scanned = False
        if file_ext in _IMAGE_EXTS:
            is_scanned = True
        elif file_ext in _FAST_PDF_EXTS and settings.ocr_enabled:
            char_threshold = page_count * settings.ocr_scanned_threshold
            if len(markdown_text.strip()) < char_threshold:
                is_scanned = True
                logger.info(
                    "Detected scanned PDF: %s (%d pages, %d chars, threshold=%d)",
                    filename, page_count, len(markdown_text.strip()), char_threshold,
                )

        logger.info(
            "Parsed %s: %d pages, %d chars, %d est. tokens, type=%s, "
            "parser=%s, scanned=%s, time=%.2fs",
            filename,
            page_count,
            len(markdown_text),
            token_estimate,
            doc_type.value,
            parser_used,
            is_scanned,
            elapsed,
        )

        return ParsedDocument(
            filename=filename,
            content=markdown_text,
            page_count=page_count,
            file_size_bytes=file_size,
            doc_type=doc_type,
            token_estimate=token_estimate,
            file_path=file_path,
            is_scanned=is_scanned,
        )

    except Exception as exc:
        error_content = f"[ERROR] Failed to parse {filename}: {exc}"
        logger.error("Error parsing %s: %s", filename, exc, exc_info=True)
        doc_type = classify_document(filename, "")
        return ParsedDocument(
            filename=filename,
            content=error_content,
            page_count=0,
            file_size_bytes=file_size,
            doc_type=doc_type,
            token_estimate=len(error_content) // 4,
            file_path=file_path,
        )


async def parse_all(
    file_paths: list[tuple[Path, str]],
    on_parsed: Optional[Callable[[ParsedDocument], None]] = None,
    max_concurrent: int | None = None,
) -> list[ParsedDocument]:
    """Parse all documents with bounded concurrency.

    Uses asyncio.Semaphore to limit parallel parsing.
    With fast parsers, concurrency limit is raised to 5.
    Calls on_parsed callback after each file for SSE streaming progress.
    """
    if not file_paths:
        return []

    if max_concurrent is None:
        # Fast parsers can handle higher concurrency
        max_concurrent = 5

    semaphore = asyncio.Semaphore(max_concurrent)

    async def _parse_one(
        index: int, file_path: Path, filename: str
    ) -> tuple[int, ParsedDocument]:
        async with semaphore:
            logger.info(
                "Parsing document %d/%d: %s",
                index + 1,
                len(file_paths),
                filename,
            )
            parsed = await parse_document(file_path, filename)
            if on_parsed is not None:
                on_parsed(parsed)
            return (index, parsed)

    tasks = [_parse_one(i, fp, fn) for i, (fp, fn) in enumerate(file_paths)]
    indexed_results = await asyncio.gather(*tasks)
    indexed_results_sorted = sorted(indexed_results, key=lambda x: x[0])
    results = [doc for _, doc in indexed_results_sorted]

    logger.info(
        "Parsing complete: %d documents, %d total tokens",
        len(results),
        sum(d.token_estimate for d in results),
    )
    return results


# ── Classification rules ──────────────────────────────────────────────────────

# Pattern → DocumentType mapping (order matters — first match wins)
_CLASSIFICATION_RULES: list[tuple[re.Pattern, DocumentType]] = [
    (re.compile(r"technin|specifikacij", re.IGNORECASE), DocumentType.TECHNICAL_SPEC),
    (re.compile(r"sutart", re.IGNORECASE), DocumentType.CONTRACT),
    (re.compile(r"kvietim|skelbim", re.IGNORECASE), DocumentType.INVITATION),
    (re.compile(r"kvalifikacij", re.IGNORECASE), DocumentType.QUALIFICATION),
    (re.compile(r"vertinim|kriterij", re.IGNORECASE), DocumentType.EVALUATION),
    (re.compile(r"pried|forma|šablon|sablon", re.IGNORECASE), DocumentType.ANNEX),
]


def classify_document(filename: str, content_preview: str) -> DocumentType:
    """Classify document type using filename and content heuristics.

    Checks filename first, then falls back to content preview.
    Uses Lithuanian keyword patterns. Case-insensitive.
    """
    # Check filename first
    for pattern, doc_type in _CLASSIFICATION_RULES:
        if pattern.search(filename):
            return doc_type

    # Fall back to content preview
    for pattern, doc_type in _CLASSIFICATION_RULES:
        if pattern.search(content_preview):
            return doc_type

    return DocumentType.OTHER


def _estimate_pages(content: str, file_ext: str) -> int:
    """Estimate page count from content length.

    - ~3000 chars per page for text documents
    - XLSX: 1 page per sheet (estimate from markdown section headers)
    - Minimum 1 page if there's any content
    """
    if not content:
        return 0

    if file_ext in (".xlsx", ".xls"):
        # Count sheet-like sections in markdown (## headers often indicate sheets)
        sheet_markers = len(re.findall(r"^##\s", content, re.MULTILINE))
        return max(sheet_markers, 1)

    # General estimate: ~3000 chars per page
    pages = max(len(content) // 3000, 1)
    return pages
